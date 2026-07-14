from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\model\CTSF")
SOURCE = ROOT / "Prompt-Z_Gate_Collapse_诊断与_Gamma-only_实验总结.md"
OUTPUT = ROOT / "Prompt-Z_Gate_Collapse_诊断与_Gamma-only_实验总结_公式版.docx"


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color="6B7280")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_font(run, size=9, color="6B7280")


def clean_inline(text):
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = re.sub(r"\$([^$]+)\$", lambda m: latex_to_linear(m.group(1)), text)
    return text


def latex_to_linear(s):
    s = s.strip()
    # Handle structures before replacing their internal commands/braces.
    s = re.sub(
        r"\\frac\{([^{}]+)\}\{([^{}]+)\}",
        lambda m: "(" + m.group(1) + ")/(" + m.group(2) + ")",
        s,
    )
    s = s.replace(r"\overline{\gamma}", "γ̄")
    s = s.replace(r"\overline{m}", "m̄")
    s = s.replace(r"\overline{\gamma}_{\mathrm{oracle}}", "γ̄_(oracle)")
    s = s.replace(r"\not\Rightarrow", "⇏")
    replacements = {
        r"\hat{Y}": "Ŷ",
        r"\gamma": "γ",
        r"\Delta": "Δ",
        r"\lambda": "λ",
        r"\approx": "≈",
        r"\qquad": "  ",
        r"\quad": " ",
        r"\cdot": "·",
        r"\times": "×",
        r"\in": "∈",
        r"\ldots": "…",
        r"\%": "%",
        r"\to": "→",
        r"\longrightarrow": "⟶",
        r"\Rightarrow": "⇒",
        r"\lVert": "‖",
        r"\rVert": "‖",
        r"\;": " ",
        r"\,": " ",
    }
    for a, b in replacements.items():
        s = s.replace(a, b)
    s = re.sub(r"\\mathrm\{([^{}]+)\}", lambda m: m.group(1).replace(r"\ ", " "), s)
    s = re.sub(r"\\text\{([^{}]+)\}", lambda m: '"' + m.group(1).replace(r"\ ", " ") + '"', s)
    s = re.sub(r"\\mathcal\{([^{}]+)\}", r"\1", s)
    s = re.sub(r"\\overline\{([^{}]+)\}", lambda m: m.group(1) + "̄", s)
    s = re.sub(
        r"\\underbrace\{([^{}]+)\}_\{([^{}]+)\}",
        lambda m: r"\underbrace(" + m.group(1) + ")_" + m.group(2),
        s,
    )
    # Word's linear equation parser does not build underbraces reliably in
    # automation, so preserve the semantic grouping with parenthesized labels.
    s = s.replace(r"\underbrace", "")
    s = s.replace("\\", " ")
    s = s.replace("{", "(").replace("}", ")")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def new_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = next(
        n for n in numbering.findall(qn("w:num"))
        if int(n.get(qn("w:numId"))) == int(style_num_id)
    )
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    used = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(used) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def assign_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = num_pr.get_or_add_ilvl()
    ilvl.val = 0
    num = num_pr.get_or_add_numId()
    num.val = num_id


def configure_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Title": (25, "111827", 0, 6),
        "Subtitle": (11, "64748B", 0, 16),
        "Heading 1": (16, "1F4D78", 16, 8),
        "Heading 2": (13, "2E5E88", 13, 6),
        "Heading 3": (11.5, "334155", 10, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.38)
        st.paragraph_format.first_line_indent = Inches(-0.18)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.2


def add_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "CBD5E1")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_equation(doc, latex):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_together = True
    run = p.add_run("[EQ] " + latex_to_linear(latex))
    set_font(run, "Cambria Math", 12, color="111827")


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "5B8DB8")
    borders.append(left)
    p_pr.append(borders)
    run = p.add_run(clean_inline(text))
    set_font(run, size=10.5, bold=True, color="264B6A")


def add_markdown_table(doc, rows):
    if len(rows) < 2:
        return
    data = [r for i, r in enumerate(rows) if i != 1]
    cols = len(data[0])
    table = doc.add_table(rows=len(data), cols=cols)
    table.autofit = False
    table.style = "Table Grid"
    usable = 6.7
    widths = [usable / cols] * cols
    if cols == 3:
        widths = [3.3, 1.55, 1.85]
    elif cols == 2:
        widths = [2.5, 4.2]
    for i, row in enumerate(data):
        for j, value in enumerate(row[:cols]):
            cell = table.cell(i, j)
            cell.width = Inches(widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(clean_inline(value.strip()))
            set_font(run, size=9.5, bold=(i == 0), color="FFFFFF" if i == 0 else "1F2937")
            if i == 0:
                shade_cell(cell, "385D7A")
            elif i % 2 == 0:
                shade_cell(cell, "F3F6F8")
    set_repeat_table_header(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]

    header = sec.header.paragraphs[0]
    header.text = "Prompt-Z · Experiment Note"
    set_font(header.runs[0], size=8.5, color="64748B")
    add_page_number(sec.footer.paragraphs[0])

    title = doc.add_paragraph(style="Title")
    title.add_run("Prompt-Z Gate Collapse")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("诊断与 Gamma-only 实验总结  ·  技术实验记录")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    add_rule(rule)

    i = 1
    in_equation = False
    eq = []
    in_number_list = False
    number_id = None
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped == "$$":
            if in_equation:
                add_equation(doc, " ".join(eq))
                eq = []
                in_equation = False
            else:
                in_equation = True
            i += 1
            continue
        if in_equation:
            eq.append(stripped)
            i += 1
            continue
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append([x.strip() for x in lines[i].strip().strip("|").split("|")])
                i += 1
            add_markdown_table(doc, table_rows)
            continue
        if stripped.startswith("### "):
            in_number_list = False
            doc.add_paragraph(clean_inline(stripped[4:]), style="Heading 3")
        elif stripped.startswith("## "):
            in_number_list = False
            doc.add_paragraph(clean_inline(stripped[3:]), style="Heading 2")
        elif stripped.startswith("# "):
            in_number_list = False
            doc.add_paragraph(clean_inline(stripped[2:]), style="Heading 1")
        elif stripped.startswith("> "):
            in_number_list = False
            add_quote(doc, stripped[2:])
        elif re.match(r"^- ", stripped):
            in_number_list = False
            doc.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            if not in_number_list:
                number_id = new_numbering_id(doc)
                in_number_list = True
            p = doc.add_paragraph(clean_inline(re.sub(r"^\d+\. ", "", stripped)), style="List Number")
            assign_numbering(p, number_id)
        else:
            in_number_list = False
            p = doc.add_paragraph()
            run = p.add_run(clean_inline(stripped))
            set_font(run, size=10.5, color="1F2937")
        i += 1

    doc.core_properties.title = "Prompt-Z Gate Collapse 诊断与 Gamma-only 实验总结"
    doc.core_properties.subject = "Prompt-Z 实验记录"
    doc.core_properties.author = ""
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
